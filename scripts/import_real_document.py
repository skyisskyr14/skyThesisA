import sys
from docx import Document
if len(sys.argv) < 2: raise SystemExit('用法: python scripts/import_real_document.py <file_path>')
doc=Document(sys.argv[1])
print({'paragraphs': len(doc.paragraphs), 'tables': len(doc.tables)})
