from __future__ import annotations

import re

from docx.document import Document as DocxDocument

from app.template_engine.utils import alignment_name, get_font_name, pt_value


def extract_caption_and_table_rules(document: DocxDocument) -> tuple[dict, dict, list[dict]]:
    figure_samples: list[str] = []
    table_samples: list[str] = []
    evidence: list[dict] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if re.match(r"图\s*\d+[-－.]\d+", text):
            figure_samples.append(text)
            evidence.append({"field": "figures.samples", "source": "图题段落", "text": text})
        if re.match(r"表\s*\d+[-－.]\d+", text):
            table_samples.append(text)
            evidence.append({"field": "tables.samples", "source": "表题段落", "text": text})

    figures = {
        "caption_position": "below",
        "caption_pattern": "图{chapter}-{index} {title}",
        "font_name": "宋体",
        "font_size_pt": 10.5,
        "alignment": "center",
        "samples": figure_samples[:5],
    }
    tables = {
        "caption_position": "above",
        "caption_pattern": "表{chapter}-{index} {title}",
        "use_three_line_table": _detect_three_line_table(document),
        "font_name": "宋体",
        "font_size_pt": 10.5,
        "alignment": "center",
        "samples": table_samples[:5],
        "table_count": len(document.tables),
        "border_summary": _border_summary(document),
        "header_row_style": _header_row_style(document),
        "cell_font": _cell_font(document),
        "cell_alignment": _cell_alignment(document),
    }
    if document.tables:
        evidence.append({"field": "tables", "source": "document.tables", "text": f"检测到 {len(document.tables)} 个表格"})
    return figures, tables, evidence


def _detect_three_line_table(document: DocxDocument) -> bool:
    if not document.tables:
        return False
    for table in document.tables:
        text = " ".join(cell.text for row in table.rows for cell in row.cells)
        if "三线表" in text:
            return True
        if len(table.rows) >= 2:
            return True
    return False


def _border_summary(document: DocxDocument) -> str:
    if not document.tables:
        return "no_table"
    return "detected_table_borders_or_grid; three_line_candidate" if _detect_three_line_table(document) else "unknown"


def _header_row_style(document: DocxDocument) -> dict:
    if not document.tables or not document.tables[0].rows:
        return {}
    row = document.tables[0].rows[0]
    return {"cell_count": len(row.cells), "texts": [cell.text for cell in row.cells]}


def _cell_font(document: DocxDocument) -> str | None:
    if not document.tables:
        return None
    for row in document.tables[0].rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.text.strip():
                        return get_font_name(run)
    return None


def _cell_alignment(document: DocxDocument) -> str | None:
    if not document.tables:
        return None
    for row in document.tables[0].rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if paragraph.text.strip():
                    return alignment_name(paragraph.alignment)
    return None
