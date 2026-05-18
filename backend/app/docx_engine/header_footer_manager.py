from __future__ import annotations

from typing import Any

from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class HeaderFooterManager:
    def __init__(self, rules: dict[str, Any], title: str) -> None:
        self.rules = rules.get("header_footer", {})
        self.title = title

    def apply(self, document: DocxDocument) -> list[str]:
        applied: list[str] = []
        for section in document.sections:
            if self.rules.get("has_header", True):
                header_text = self.rules.get("header_text") or self.title
                section.header.paragraphs[0].text = header_text
                applied.append("header_text")
            if self.rules.get("has_footer", True):
                footer = section.footer.paragraphs[0]
                footer.text = self.rules.get("footer_text") or ""
                if self.rules.get("has_page_number", True):
                    self._append_page_number(footer)
                    applied.append("page_number_field")
        return applied

    def _append_page_number(self, paragraph) -> None:
        paragraph.add_run(" 第 ")
        run = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)
        paragraph.add_run(" 页")
