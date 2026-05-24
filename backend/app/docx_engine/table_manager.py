from __future__ import annotations

from typing import Any

from docx.document import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from app.docx_engine.common import ALIGNMENT_MAP


class TableManager:
    def __init__(self, rules: dict[str, Any]) -> None:
        self.rules = rules.get("tables", {})

    def add_table(self, document: DocxDocument, chapter_no: int, index: int, title: str, columns: list[str], rows: list[list[str]]) -> None:
        caption_text = self._caption(chapter_no, index, title)
        if self.rules.get("caption_position", "above") == "above":
            self._add_caption(document, caption_text)
        table = document.add_table(rows=len(rows) + 1, cols=len(columns))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for col_index, column in enumerate(columns):
            table.cell(0, col_index).text = column
        for row_index, row in enumerate(rows, start=1):
            for col_index, value in enumerate(row):
                table.cell(row_index, col_index).text = value
        self._style_cells(table)
        if self.rules.get("use_three_line_table", True):
            self._apply_three_line_borders(table)
        if self.rules.get("caption_position", "above") == "below":
            self._add_caption(document, caption_text)

    def _caption(self, chapter_no: int, index: int, title: str) -> str:
        pattern = self.rules.get("caption_pattern", "表{chapter}-{index} {title}")
        return pattern.replace("{chapter}", str(chapter_no)).replace("{index}", str(index)).replace("{title}", title)

    def _add_caption(self, document: DocxDocument, text: str) -> None:
        p = document.add_paragraph(text)
        p.alignment = ALIGNMENT_MAP.get(self.rules.get("alignment", "center"))
        for run in p.runs:
            run.font.name = self.rules.get("font_name", "宋体")
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.rules.get("font_name", "宋体"))
            run.font.size = Pt(float(self.rules.get("font_size_pt", 10.5)))

    def _style_cells(self, table) -> None:
        alignment = ALIGNMENT_MAP.get(self.rules.get("cell_alignment") or self.rules.get("alignment", "center"))
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = alignment
                    for run in p.runs:
                        run.font.name = self.rules.get("font_name", "宋体")
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.rules.get("font_name", "宋体"))
                        run.font.size = Pt(float(self.rules.get("font_size_pt", 10.5)))

    def _apply_three_line_borders(self, table) -> None:
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                borders = OxmlElement("w:tcBorders")
                for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                    border = OxmlElement(f"w:{edge}")
                    border.set(qn("w:val"), "nil")
                    if edge == "top" and row_index == 0:
                        border.set(qn("w:val"), "single")
                        border.set(qn("w:sz"), "12")
                    if edge == "bottom" and row_index in {0, len(table.rows) - 1}:
                        border.set(qn("w:val"), "single")
                        border.set(qn("w:sz"), "12")
                    borders.append(border)
                tc_pr.append(borders)
