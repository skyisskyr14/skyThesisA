from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class DocxEngine:
    """最小 DOCX 精排引擎示例，支持初步套用 v0.3 template_rules。"""

    def __init__(self, output_dir: Path) -> None:
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_sample(self, project_id: int, title: str, template_rules: dict[str, Any] | None = None) -> Path:
        rules = template_rules or {}
        document = Document()
        self._apply_page_rules(document, rules.get("page", {}))
        self._setup_styles(document, rules.get("body", {}))
        self._add_title(document, title, rules)
        self._add_abstract(document)
        self._add_chapter(document, title, rules)
        self._add_references(document)

        output = self.output_dir / f"project_{project_id}_sample.docx"
        document.save(output)
        return output

    def _apply_page_rules(self, document: Document, page_rules: dict[str, Any]) -> None:
        for section in document.sections:
            if page_rules.get("margin_top_cm"):
                section.top_margin = Cm(float(page_rules["margin_top_cm"]))
            if page_rules.get("margin_bottom_cm"):
                section.bottom_margin = Cm(float(page_rules["margin_bottom_cm"]))
            if page_rules.get("margin_left_cm"):
                section.left_margin = Cm(float(page_rules["margin_left_cm"]))
            if page_rules.get("margin_right_cm"):
                section.right_margin = Cm(float(page_rules["margin_right_cm"]))

    def _add_title(self, document: Document, title: str, rules: dict[str, Any]) -> None:
        heading_rule = self._heading_rule(rules, level=1)
        heading = document.add_paragraph()
        heading.alignment = ALIGNMENT_MAP.get(heading_rule.get("alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER)
        run = heading.add_run(title)
        run.bold = bool(heading_rule.get("bold", True))
        run.font.name = heading_rule.get("font_name") or "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), heading_rule.get("font_name") or "黑体")
        run.font.size = Pt(float(heading_rule.get("font_size_pt") or 18))

    def _add_abstract(self, document: Document) -> None:
        document.add_heading("摘  要", level=1)
        document.add_paragraph("摘要占位：本论文围绕系统背景、设计实现、测试结果和应用价值展开说明。")
        document.add_paragraph("关键词：Thesis Agent；DOCX 精排；MemoryGuard；多 Agent")

    def _add_chapter(self, document: Document, title: str, rules: dict[str, Any]) -> None:
        heading_rule = self._heading_rule(rules, level=1)
        heading = document.add_heading("第1章 绪论", level=1)
        heading.alignment = ALIGNMENT_MAP.get(heading_rule.get("alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER)
        for run in heading.runs:
            run.font.name = heading_rule.get("font_name") or "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), heading_rule.get("font_name") or "黑体")
            run.font.size = Pt(float(heading_rule.get("font_size_pt") or 16))
        document.add_paragraph(
            f"本文以《{title}》为研究对象，MVP 阶段先验证项目创建、模板分析、大纲生成、章节写作、DOCX 生成和最终审查闭环。"
        )
        self._add_figure_caption(document, rules.get("figures", {}))
        self._add_three_line_table(document, rules.get("tables", {}))

    def _add_figure_caption(self, document: Document, figure_rules: dict[str, Any]) -> None:
        caption = document.add_paragraph("图 1-1 系统架构图占位：白底黑字、线条不交叉、节点不重叠。")
        caption.alignment = ALIGNMENT_MAP.get(figure_rules.get("alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER)
        for run in caption.runs:
            run.font.name = figure_rules.get("font_name") or "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), figure_rules.get("font_name") or "宋体")
            run.font.size = Pt(float(figure_rules.get("font_size_pt") or 10.5))

    def _add_three_line_table(self, document: Document, table_rules: dict[str, Any]) -> None:
        caption = document.add_paragraph("表 1-1 MVP 三线表示例")
        caption.alignment = ALIGNMENT_MAP.get(table_rules.get("alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER)
        table = document.add_table(rows=4, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        headers = ["模块", "职责", "后续扩展"]
        for index, text in enumerate(headers):
            table.cell(0, index).text = text
        rows = [
            ["TemplateAnalyzerAgent", "模板分析", "接入真实 DOCX 样式提取"],
            ["DocxEngine", "精准排版", "支持局部替换和格式保护"],
            ["MemoryGuard", "记忆纠错", "阻止历史错误重复导出"],
        ]
        for row_index, row in enumerate(rows, start=1):
            for col_index, text in enumerate(row):
                table.cell(row_index, col_index).text = text
        if table_rules.get("use_three_line_table", True):
            self._apply_three_line_table(table)

    def _add_references(self, document: Document) -> None:
        document.add_heading("参考文献", level=1)
        document.add_paragraph("[1] 后续由 CitationAgent 根据引用位置自动生成。")

    def _setup_styles(self, document: Document, body_rules: dict[str, Any]) -> None:
        normal = document.styles["Normal"]
        font_name = body_rules.get("font_cn") or "宋体"
        normal.font.name = font_name
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        normal.font.size = Pt(float(body_rules.get("font_size_pt") or 12))
        paragraph_format = normal.paragraph_format
        if body_rules.get("line_spacing"):
            paragraph_format.line_spacing = float(body_rules["line_spacing"])
        if body_rules.get("first_line_indent_chars"):
            paragraph_format.first_line_indent = Pt(float(body_rules["first_line_indent_chars"]) * float(body_rules.get("font_size_pt") or 12))

    def _apply_three_line_table(self, table) -> None:
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                borders = OxmlElement("w:tcBorders")
                top = OxmlElement("w:top")
                bottom = OxmlElement("w:bottom")
                for border in [top, bottom]:
                    border.set(qn("w:val"), "nil")
                if row_index in {0, len(table.rows) - 1}:
                    bottom.set(qn("w:val"), "single")
                    bottom.set(qn("w:sz"), "12")
                if row_index == 0:
                    top.set(qn("w:val"), "single")
                    top.set(qn("w:sz"), "12")
                borders.append(top)
                borders.append(bottom)
                tc_pr.append(borders)

    def _heading_rule(self, rules: dict[str, Any], level: int) -> dict[str, Any]:
        for heading in rules.get("headings", []):
            if heading.get("level") == level:
                return heading
        return {"font_name": "黑体", "font_size_pt": 16, "alignment": "center", "bold": True}
