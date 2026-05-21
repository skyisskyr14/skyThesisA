from __future__ import annotations

from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.docx_engine.docx_models import PaperMeta


class CoverManager:
    def add_cover(self, document: DocxDocument, meta: PaperMeta) -> None:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(meta.title)
        run.bold = True
        run.font.size = Pt(22)
        document.add_paragraph("\n\n")
        for label, value in [
            ("学校", meta.school),
            ("专业", meta.major),
            ("班级", meta.class_name),
            ("学生姓名", meta.author),
            ("学号", meta.student_no),
            ("指导教师", meta.teacher),
            ("日期", meta.date),
        ]:
            p = document.add_paragraph(f"{label}：{value}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_page_break()
