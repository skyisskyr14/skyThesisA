import hashlib
import uuid
from pathlib import Path
from docx import Document
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from sqlalchemy.orm import Session

from app.database import BASE_DIR, get_db
from app.models.entities import PaperFile, ThesisProject
from app.schemas.common import FileRead

router = APIRouter(prefix="/api/thesis/files", tags=["files"])
UPLOAD_DIR = BASE_DIR / "storage" / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
ALLOWED = {'.docx', '.txt', '.md'}

@router.post('/upload', response_model=FileRead)
def upload_file(project_id: int = Form(...), file_role: str = Form('unknown'), file: UploadFile = File(...), db: Session = Depends(get_db)):
    project = db.get(ThesisProject, project_id)
    if not project: raise HTTPException(404, '项目不存在')
    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED: raise HTTPException(400, '仅支持 .docx .txt .md')
    content = file.file.read()
    checksum = hashlib.sha256(content).hexdigest()
    pdir = UPLOAD_DIR / str(project_id); pdir.mkdir(parents=True, exist_ok=True)
    stored = f"{uuid.uuid4().hex}{ext}"
    path = pdir / stored
    path.write_bytes(content)
    row = PaperFile(project_id=project_id, filename=file.filename, file_type=ext.lstrip('.'), purpose=file_role, path=str(path), status='completed')
    db.add(row); db.commit(); db.refresh(row)
    return row

@router.get('/{project_id}', response_model=list[FileRead])
def list_files(project_id: int, db: Session = Depends(get_db)):
    return db.query(PaperFile).filter(PaperFile.project_id == project_id).order_by(PaperFile.created_at.desc()).all()

@router.post('/{file_id}/identify')
def identify(file_id: int, db: Session = Depends(get_db)):
    row = db.get(PaperFile, file_id)
    if not row: raise HTTPException(404, 'file not found')
    name = row.filename.lower()
    role = 'unknown'
    if '模板' in name: role = 'thesis_template'
    elif '开题' in name: role = 'proposal_report'
    elif '任务书' in name: role = 'task_book'
    elif '初稿' in name: role = 'existing_draft'
    row.purpose = role
    db.commit()
    return {'file_id': file_id, 'file_role': role}

@router.post('/{file_id}/parse')
def parse_file(file_id: int, db: Session = Depends(get_db)):
    row = db.get(PaperFile, file_id)
    if not row: raise HTTPException(404, 'file not found')
    p = Path(row.path)
    if not p.exists(): raise HTTPException(404, 'stored file not found')
    if p.suffix.lower() == '.docx':
        doc = Document(str(p))
        return {'file_id': file_id, 'paragraph_count': len(doc.paragraphs), 'table_count': len(doc.tables)}
    text = p.read_text(encoding='utf-8', errors='ignore')
    return {'file_id': file_id, 'chars': len(text)}

@router.delete('/{file_id}')
def delete_file(file_id: int, db: Session = Depends(get_db)):
    row = db.get(PaperFile, file_id)
    if not row: raise HTTPException(404, 'file not found')
    p = Path(row.path)
    if p.exists(): p.unlink()
    db.delete(row); db.commit()
    return {'ok': True}
