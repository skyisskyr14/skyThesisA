from pathlib import Path
from docx import Document
from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session

from app.database import get_db
from app.models.entities import PaperFile, PaperVersion, ThesisProject

router = APIRouter(prefix='/api/thesis/papers', tags=['papers'])

@router.post('/import-docx')
def import_docx(payload: dict, db: Session = Depends(get_db)):
    project_id = int(payload.get('project_id', 0)); file_id = int(payload.get('file_id', 0))
    p = db.get(ThesisProject, project_id)
    f = db.get(PaperFile, file_id)
    if not p or not f: raise HTTPException(404, '项目或文件不存在')
    if Path(f.path).suffix.lower() != '.docx': raise HTTPException(400, '仅支持 DOCX 导入')
    doc = Document(f.path)
    paras = [x.text.strip() for x in doc.paragraphs if x.text.strip()]
    title = paras[0] if paras else p.title
    chapters=[]; ch_no=1; sec_no=1; bi=1
    cur={"chapter_id":f"ch_{ch_no}","chapter_no":ch_no,"title":"导入章节","sections":[{"section_id":f"ch_{ch_no}_sec_{sec_no}","section_no":f"{ch_no}.{sec_no}","level":2,"title":"正文","blocks":[]}]} 
    for i,t in enumerate(paras[:200],start=1):
        cur['sections'][0]['blocks'].append({"block_id":f"ch_{ch_no}_sec_{sec_no}_p_{bi}","block_type":"paragraph","text":t,"source_location":{"file_id":file_id,"paragraph_index":i}}); bi+=1
    chapters.append(cur)
    paper_doc={"project_id":project_id,"source_file_id":file_id,"meta":{"title":title},"abstract":{"zh":[],"zh_keywords":[],"en":[],"en_keywords":[]},"chapters":chapters,"references":[],"import_warnings":[]}
    vcount = db.query(PaperVersion).filter(PaperVersion.project_id==project_id).count()+1
    v=PaperVersion(project_id=project_id,version_no=f"v{vcount}",version_name='导入原始论文',version_type='import_original',source_file_id=file_id,paper_document_json=paper_doc,change_summary='从真实DOCX导入论文内容',total_tokens_used=0,is_current=True)
    db.query(PaperVersion).filter(PaperVersion.project_id==project_id).update({PaperVersion.is_current: False})
    db.add(v); db.commit(); db.refresh(v)
    return {"success":True,"project_id":project_id,"file_id":file_id,"version_id":v.id,"detected_title":title,"chapter_count":len(chapters),"section_count":1,"paragraph_block_count":len(chapters[0]['sections'][0]['blocks']),"table_count":len(doc.tables),"figure_caption_count":0,"reference_count":0,"warnings":[],"structure_preview":paper_doc}
